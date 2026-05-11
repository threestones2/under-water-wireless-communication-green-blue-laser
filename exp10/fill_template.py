"""Fill conference template with title, abstract, and keywords."""
from docx import Document
from docx.shared import Pt
import copy
import re

TEMPLATE = r"D:\509所-余重俊（公开）\8.蓝绿课题\小论文\会议论文\conference-template-a4.docx"
OUTPUT   = r"D:\509所-余重俊（公开）\8.蓝绿课题\小论文\会议论文\APCCAS_SA-MCS_Paper.docx"

EM_DASH = "—"

TITLE = "A Computationally Efficient Semi-Analytical Monte Carlo Method for Real-Time Underwater Optical Channel Estimation"

ABSTRACT_BODY = (
    "Real-time channel estimation is critical for adaptive underwater wireless "
    "optical communication (UWOC) systems deployed on autonomous underwater "
    "vehicles (AUVs), yet conventional Monte Carlo (MC) methods impose prohibitive "
    "computational costs. This paper presents a Semi-Analytical Monte Carlo "
    "Simulation (SA-MCS) framework that replaces physical photon reception with "
    "closed-form local estimation at each scattering node, eliminating the need "
    "for ray-marching and geometric hit-testing. Two scattering-angle sampling "
    "strategies are investigated: a look-up table (LUT) approach achieving O(1) "
    "constant-time sampling, and an inverse-transform method offering greater "
    "flexibility. The proposed methods are benchmarked against wavefront-coupled "
    "importance-sampling MC (WCI-MC) and ground-truth physical MC across three "
    "water types including clear ocean, coastal ocean, and turbid harbor, at "
    "photon counts from 10^3 to 10^5. Experimental results demonstrate that "
    "SA-MCS (LUT) achieves path loss estimation accuracy within 0.5 dB of the "
    "reference while reducing computation time by over an order of magnitude "
    "compared to WCI-MC. The LUT-based architecture, with its fixed memory "
    "footprint and constant-time sampling, is particularly well-suited for "
    "hardware acceleration on resource-constrained embedded platforms, enabling "
    "real-time channel adaptation for next-generation intelligent underwater "
    "circuits and systems."
)

KEYWORDS_BODY = (
    "underwater wireless optical communication, Monte Carlo simulation, "
    "semi-analytical method, channel estimation, edge computing, embedded systems"
)


def replace_paragraph_keeping_first_run(doc, old_prefix, new_prefix, new_body):
    """Find paragraph starting with old_prefix, replace while keeping first-run formatting."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith(old_prefix):
            # Clear all runs and rebuild
            for run in para.runs:
                run.text = ""
            # Re-use first run for the prefix (preserves style: bold+italic)
            if para.runs:
                para.runs[0].text = new_prefix + new_body
            else:
                para.add_run(new_prefix + new_body)
            return True
    return False


def main():
    doc = Document(TEMPLATE)

    # === Replace Title ===
    for para in doc.paragraphs:
        if "Paper Title" in para.text:
            # Keep style, replace text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = TITLE
            else:
                run = para.add_run(TITLE)
                run.font.size = Pt(22)
            print(f"Title replaced. Style: {para.style.name}")
            break

    # === Replace Abstract ===
    abstract_prefix = f"Abstract{EM_DASH}"
    found_abs = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith(abstract_prefix):
            # Clear all runs
            for run in para.runs:
                run.text = ""
            # Add formatted prefix + body
            if para.runs:
                # Use first run for prefix with original formatting
                para.runs[0].text = abstract_prefix + ABSTRACT_BODY
            else:
                para.add_run(abstract_prefix + ABSTRACT_BODY)
            found_abs = True
            print(f"Abstract replaced. Style: {para.style.name}")
            break

    if not found_abs:
        print("WARNING: Abstract paragraph not found!")

    # === Replace Keywords ===
    keywords_prefix = f"Keywords{EM_DASH}"
    found_kw = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith(keywords_prefix):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = keywords_prefix + KEYWORDS_BODY
            else:
                para.add_run(keywords_prefix + KEYWORDS_BODY)
            found_kw = True
            print(f"Keywords replaced. Style: {para.style.name}")
            break

    if not found_kw:
        print("WARNING: Keywords paragraph not found!")

    # === Remove template guide paragraphs ===
    # Delete "Introduction (Heading 1)" and everything after it
    # Or keep it for now — user can delete later

    doc.save(OUTPUT)
    print(f"\nSaved to: {OUTPUT}")


if __name__ == "__main__":
    main()
